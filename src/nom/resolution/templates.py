"""Template storage for Resolution documents."""
import json
import os
import time
import typing
from pathlib import Path
from typing import List, Optional
from nom.resolution.schema import NghiQuyetBase

TEMPLATES_DIR = Path.home() / ".nom" / "resolution_templates"
DOCX_TEMPLATES_DIR = Path.home() / ".nom" / "resolution_docx_templates"

def _ensure_dir():
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    DOCX_TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

# --- Data Templates (JSON) ---

def save_template(name: str, data: typing.Union[NghiQuyetBase, dict]) -> str:
    """Save a document as a reusable JSON template."""
    _ensure_dir()
    filename = name.replace(" ", "_").replace("/", "-") + ".json"
    filepath = TEMPLATES_DIR / filename
    
    dict_data = data.model_dump() if isinstance(data, NghiQuyetBase) else data
    
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(dict_data, f, ensure_ascii=False, indent=2)
        
    return filename

def list_templates() -> List[dict]:
    """List all saved JSON templates."""
    _ensure_dir()
    templates = []
    for fp in sorted(TEMPLATES_DIR.glob("*.json")):
        try:
            with open(fp, "r", encoding="utf-8") as f:
                data = json.load(f)
            templates.append({
                "filename": fp.name,
                "name": fp.stem.replace("_", " "),
                "trich_yeu": data.get("trich_yeu", ""),
                "co_quan_ban_hanh": data.get("co_quan_ban_hanh", ""),
            })
        except Exception:
            pass
    return templates

def load_template(filename: str) -> Optional[dict]:
    """Load a JSON template by filename."""
    filepath = TEMPLATES_DIR / filename
    if not filepath.exists():
        return None
    with open(filepath, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data

def delete_template(filename: str) -> bool:
    """Delete a JSON template by filename."""
    filepath = TEMPLATES_DIR / filename
    if filepath.exists():
        filepath.unlink()
        return True
    return False

import re
import unicodedata
import io

def remove_vietnamese_accents(s: str) -> str:
    s = re.sub(r'[àáạảãâầấậẩẫăằắặẳẵ]', 'a', s)
    s = re.sub(r'[èéẹẻẽêềếệểễ]', 'e', s)
    s = re.sub(r'[ìíịỉĩ]', 'i', s)
    s = re.sub(r'[òóọỏõôồốộổỗơờớợởỡ]', 'o', s)
    s = re.sub(r'[ùúụủũưừứựửữ]', 'u', s)
    s = re.sub(r'[ỳýỵỷỹ]', 'y', s)
    s = re.sub(r'[đ]', 'd', s)
    s = re.sub(r'[ÀÁẠẢÃÂẦẤẬẨẪĂẰẮẶẲẴ]', 'A', s)
    s = re.sub(r'[ÈÉẸẺẼÊỀẾỆỂỄ]', 'E', s)
    s = re.sub(r'[ÌÍỊỈĨ]', 'I', s)
    s = re.sub(r'[ÒÓỌỎÕÔỒỐỘỔỖƠỜỚỢỞỠ]', 'O', s)
    s = re.sub(r'[ÙÚỤỦŨƯỪỨỰỬỮ]', 'U', s)
    s = re.sub(r'[ỲÝỴỶỸ]', 'Y', s)
    s = re.sub(r'[Đ]', 'D', s)
    # Remove accents using unicodedata just in case
    s = unicodedata.normalize('NFKD', s).encode('ASCII', 'ignore').decode('utf-8')
    return s

def generate_var_name(text_before: str) -> str:
    # Clean up string
    clean_text = remove_vietnamese_accents(text_before.strip())
    # Remove non-alphanumeric chars (keep spaces temporarily)
    clean_text = re.sub(r'[^a-zA-Z0-9\s]', '', clean_text)
    words = clean_text.strip().split()
    if not words:
        import uuid
        return f"var_{uuid.uuid4().hex[:6]}"
    # Get up to last 4 words for variable name
    var_words = words[-4:]
    var_name = "_".join(var_words).lower()
    return var_name

def auto_fill_template_variables(content: bytes) -> bytes:
    try:
        from docx import Document
        buffer = io.BytesIO(content)
        doc = Document(buffer)
        
        # Regex to match 4 or more dots or underscores
        blank_pattern = re.compile(r'([\.\_]{4,})')
        
        # Iterate over all paragraphs (in document body and tables)
        paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
                    
        modified = False
        
        for p in paragraphs:
            # Check if paragraph text contains blanks
            if blank_pattern.search(p.text):
                # We will rewrite the runs in this paragraph
                # But it's easier to just reconstruct the whole text and assign it to p.text
                # Note: This might lose formatting within the paragraph, but it's acceptable for simple templates
                new_text = ""
                parts = blank_pattern.split(p.text)
                # parts will look like: ["Họ và tên: ", ".......", " Năm sinh: ", "_____", ""]
                for i in range(len(parts)):
                    if blank_pattern.match(parts[i]):
                        # It's a blank! Let's generate a variable name based on accumulated text so far
                        var_name = generate_var_name(new_text)
                        new_text += f"{{{{ {var_name} }}}}"
                        modified = True
                    else:
                        new_text += parts[i]
                
                # Assign back to paragraph (clears formatting, uses paragraph default)
                if modified:
                    p.text = new_text

        if modified:
            out_buffer = io.BytesIO()
            doc.save(out_buffer)
            return out_buffer.getvalue()
        
        return content
    except Exception as e:
        print(f"Error in auto_fill_template_variables: {e}")
        return content

# --- Layout Templates (.docx) ---

def save_docx_template(filename: str, content: bytes) -> str:
    """Save an uploaded .docx layout template."""
    _ensure_dir()
    
    # Auto-convert dots to template variables
    content = auto_fill_template_variables(content)
    
    # Sanitize filename
    safe_name = os.path.basename(filename).replace(" ", "_")
    if not safe_name.endswith(".docx"):
        safe_name += ".docx"
    filepath = DOCX_TEMPLATES_DIR / safe_name
    with open(filepath, "wb") as f:
        f.write(content)
    return safe_name

def list_docx_templates() -> List[dict]:
    """List all saved .docx templates."""
    _ensure_dir()
    templates = []
    for fp in sorted(DOCX_TEMPLATES_DIR.glob("*.docx")):
        templates.append({
            "filename": fp.name,
            "name": fp.stem.replace("_", " "),
        })
    return templates

def delete_docx_template(filename: str) -> bool:
    """Delete a .docx template by filename."""
    filepath = DOCX_TEMPLATES_DIR / filename
    if filepath.exists():
        filepath.unlink()
        return True
    return False

def get_docx_template_path(filename: str) -> Optional[str]:
    """Get absolute path to a .docx template."""
    filepath = DOCX_TEMPLATES_DIR / filename
    if filepath.exists():
        return str(filepath)
    return None

def extract_docx_variables(filename: str) -> List[str]:
    """Extract all Jinja2 variables {{ var }} from a .docx template."""
    filepath = get_docx_template_path(filename)
    if not filepath:
        return []
    try:
        from docxtpl import DocxTemplate
        doc = DocxTemplate(filepath)
        vars_set = doc.get_undeclared_template_variables()
        return sorted(list(vars_set))
    except Exception as e:
        print(f"Error extracting variables: {e}")
        return []


