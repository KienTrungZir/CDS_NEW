import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 4, "val": "single", "color": "000000"},
        bottom={"sz": 4, "val": "single", "color": "000000"},
        left={"sz": 4, "val": "single", "color": "000000"},
        right={"sz": 4, "val": "single", "color": "000000"},
    )
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

class DynamicWordExporter:
    def export(self, data: dict) -> io.BytesIO:
        if "header" in data or "title_and_bases" in data:
            return self.export_nd30(data)
            
        doc = Document()
        
        # Cài đặt font mặc định
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        
        blocks = data.get("blocks", [])
        
        for block in blocks:
            b_type = block.get("type", "paragraph")
            bold = block.get("bold", False)
            italic = block.get("italic", False)
            font_size = block.get("font_size")
            
            if b_type == "header_split":
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                table.columns[0].width = Inches(3)
                table.columns[1].width = Inches(3.5)
                
                # Cột trái
                c1 = table.cell(0, 0).paragraphs[0]
                c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                left_text = block.get("left", "")
                if left_text:
                    c1.add_run(left_text).bold = True
                
                # Cột phải
                c2 = table.cell(0, 1).paragraphs[0]
                c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                right_text = block.get("right", "")
                if right_text:
                    lines = right_text.split('\n')
                    for i, line in enumerate(lines):
                        run = c2.add_run(line + ("\n" if i < len(lines)-1 else ""))
                        run.bold = True
            
            elif b_type == "title":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(block.get("text", "").upper())
                run.bold = True
                run.font.size = Pt(font_size or 14)
                
            elif b_type == "paragraph":
                p = doc.add_paragraph()
                text = block.get("text", "")
                run = p.add_run(text)
                run.bold = bold
                run.italic = italic
                if font_size:
                    run.font.size = Pt(font_size)
                    
                align_str = block.get("align", "left")
                if align_str == "center":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align_str == "right":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif align_str == "justify":
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            elif b_type == "list_item":
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run(f"- {block.get('text', '')}")
                run.bold = bold
                run.italic = italic
                if font_size:
                    run.font.size = Pt(font_size)

            elif b_type == "table":
                headers = block.get("headers", [])
                rows = block.get("rows", [])
                cols_cnt = max(len(headers), max([len(r) for r in rows] or [1]))
                if cols_cnt > 0:
                    tbl = doc.add_table(rows=0, cols=cols_cnt)
                    tbl.style = 'Table Grid'
                    
                    # Header row
                    if headers:
                        hdr_cells = tbl.add_row().cells
                        for idx, h in enumerate(headers):
                            if idx < cols_cnt:
                                p = hdr_cells[idx].paragraphs[0]
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                r = p.add_run(str(h))
                                r.bold = True
                                r.font.size = Pt(12)
                                set_cell_border(hdr_cells[idx], top={"sz": 6, "val": "single"}, bottom={"sz": 6, "val": "single"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"})
                    
                    # Data rows
                    for r_data in rows:
                        row_cells = tbl.add_row().cells
                        for idx, val in enumerate(r_data):
                            if idx < cols_cnt:
                                p = row_cells[idx].paragraphs[0]
                                r = p.add_run(str(val))
                                r.font.size = Pt(12)
                                set_cell_border(row_cells[idx], top={"sz": 4, "val": "single"}, bottom={"sz": 4, "val": "single"}, left={"sz": 4, "val": "single"}, right={"sz": 4, "val": "single"})
                    doc.add_paragraph()

            elif b_type == "divider":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run("________________________________________")
                r.font.size = Pt(10)

            elif b_type == "signature_split":
                doc.add_paragraph("\n")
                table = doc.add_table(rows=1, cols=2)
                
                # Cột trái (Nơi nhận)
                c1 = table.cell(0, 0).paragraphs[0]
                run1 = c1.add_run("Nơi nhận:\n")
                run1.bold = True
                run1.italic = True
                run1.font.size = Pt(12)
                left_text = block.get("left", "")
                if left_text:
                    for line in left_text.split('\n'):
                        c1.add_run(f"- {line};\n").font.size = Pt(11)
                        
                # Cột phải (Chữ ký)
                c2 = table.cell(0, 1).paragraphs[0]
                c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                right_text = block.get("right", "")
                if right_text:
                    parts = right_text.split('\n')
                    c2.add_run(parts[0]).bold = True  # Chức vụ
                    c2.add_run("\n\n\n\n")
                    if len(parts) > 1:
                        c2.add_run(parts[-1]).bold = True  # Tên
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def export_nd30(self, data: dict) -> io.BytesIO:
        doc = Document()
        
        # Cài đặt lề trang chuẩn Phụ lục I - Nghị định 30/2020/NĐ-CP (Hình 1)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.79)     # 20 mm (20 - 25 mm)
            section.bottom_margin = Inches(0.79)  # 20 mm (20 - 25 mm)
            section.left_margin = Inches(1.18)    # 30 mm (30 - 35 mm)
            section.right_margin = Inches(0.59)   # 15 mm (15 - 20 mm)
            
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        
        header = data.get("header", {})
        title_bases = data.get("title_and_bases", {})
        body = data.get("body", {})
        footer = data.get("footer", {})
        
        # 1. Header (2 cột)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(3.0)
        table.columns[1].width = Inches(3.5)
        
        # Cột trái (Cơ quan ban hành) - Hình 3: Cỡ 12-13, Đứng, Đậm
        c1 = table.cell(0, 0).paragraphs[0]
        c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if header.get("issuing_body_parent"):
            r_parent = c1.add_run(header.get("issuing_body_parent", "").upper() + "\n")
            r_parent.font.size = Pt(12)
            r_parent.bold = False
            
        r_body = c1.add_run(header.get("issuing_body", "").upper() + "\n")
        r_body.bold = True
        r_body.font.size = Pt(12)
        
        # Đường kẻ bên dưới tên cơ quan ban hành
        r_line1 = c1.add_run("__________\n")
        r_line1.font.size = Pt(10)
        
        if header.get("document_number"):
            r_num = c1.add_run(f"Số: {header.get('document_number')}")
            r_num.font.size = Pt(13)

        # Cột phải (Quốc hiệu Tiêu ngữ & Địa danh Ngày tháng) - Hình 3: Cỡ 12 (Quốc hiệu) / 13 (Tiêu ngữ) / 13-14 Nghiêng (Ngày tháng)
        c2 = table.cell(0, 1).paragraphs[0]
        c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_nat = c2.add_run(header.get("national_motto", "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM").upper() + "\n")
        r_nat.bold = True
        r_nat.font.size = Pt(12)
        r_mot = c2.add_run(header.get("motto", "Độc lập - Tự do - Hạnh phúc") + "\n")
        r_mot.bold = True
        r_mot.font.size = Pt(13)
        
        # Đường kẻ bên dưới tiêu ngữ
        r_line2 = c2.add_run("________________\n")
        r_line2.font.size = Pt(10)
        
        loc = header.get("location", "Hà Nội")
        dt = header.get("date") or {}
        day = dt.get("day", "..") if isinstance(dt, dict) else ".."
        month = dt.get("month", "..") if isinstance(dt, dict) else ".."
        year = dt.get("year", "....") if isinstance(dt, dict) else "...."
        r_date = c2.add_run(f"{loc}, ngày {day} tháng {month} năm {year}")
        r_date.italic = True
        r_date.font.size = Pt(13)
        
        doc.add_paragraph()
        
        # 2. Tên loại văn bản & Trích yếu - Hình 3: Cỡ 13-14, Đứng, Đậm
        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_name = p_name.add_run(title_bases.get("document_name", "NGHỊ ĐỊNH").upper())
        r_name.bold = True
        r_name.font.size = Pt(14)
        
        if title_bases.get("subject"):
            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_sub = p_sub.add_run(title_bases.get("subject"))
            r_sub.bold = True
            r_sub.font.size = Pt(14)
            
        # Căn cứ pháp lý - Cỡ 14, Nghiêng
        for base in title_bases.get("legal_bases", []):
            p_base = doc.add_paragraph()
            p_base.paragraph_format.left_indent = Inches(0.5)
            r_b = p_base.add_run(base if base.endswith(";") or base.endswith(".") else f"{base};")
            r_b.italic = True
            r_b.font.size = Pt(14)
            
        if title_bases.get("promulgation_statement"):
            p_stmt = doc.add_paragraph()
            p_stmt.paragraph_format.left_indent = Inches(0.5)
            r_stmt = p_stmt.add_run(title_bases.get("promulgation_statement"))
            r_stmt.font.size = Pt(14)
            
        doc.add_paragraph()
        
        # 3. Thân văn bản (Chapters / Articles / Clauses / Points) - Hình 4
        chapters = body.get("chapters", [])
        if chapters:
            for ch in chapters:
                p_ch = doc.add_paragraph()
                p_ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_ch = p_ch.add_run(f"Chương {ch.get('chapter_number', '')}\n{ch.get('chapter_title', '').upper()}")
                r_ch.bold = True
                r_ch.font.size = Pt(14)
                
                for art in ch.get("articles", []):
                    p_art = doc.add_paragraph()
                    art_head = f"Điều {art.get('article_number', '')}."
                    if art.get("article_title"):
                        art_head += f" {art.get('article_title')}"
                    r_art = p_art.add_run(art_head)
                    r_art.bold = True
                    r_art.font.size = Pt(14)
                    
                    if art.get("content"):
                        p_ac = doc.add_paragraph(art.get("content"))
                        p_ac.paragraph_format.left_indent = Inches(0.5)
                        
                    for cl in art.get("clauses", []):
                        p_cl = doc.add_paragraph(f"{cl.get('clause_number', '')}. {cl.get('content', '')}")
                        p_cl.paragraph_format.left_indent = Inches(0.5)
                        
                        for pt in cl.get("points", []):
                            p_pt = doc.add_paragraph(f"{pt.get('point_letter', '')}) {pt.get('content', '')}")
                            p_pt.paragraph_format.left_indent = Inches(0.8)
        else:
            articles = body.get("articles", [])
            for art in articles:
                p_art = doc.add_paragraph()
                art_head = f"Điều {art.get('article_number', '')}."
                if art.get("article_title"):
                    art_head += f" {art.get('article_title')}"
                r_art = p_art.add_run(art_head)
                r_art.bold = True
                r_art.font.size = Pt(14)
                
                if art.get("content"):
                    p_ac = doc.add_paragraph(art.get("content"))
                    p_ac.paragraph_format.left_indent = Inches(0.5)
                    
                for cl in art.get("clauses", []):
                    p_cl = doc.add_paragraph(f"{cl.get('clause_number', '')}. {cl.get('content', '')}")
                    p_cl.paragraph_format.left_indent = Inches(0.5)

        doc.add_paragraph("\n")
        
        # 4. Footer (Nơi nhận & Chữ ký) - Hình 4 & 5
        tbl_ft = doc.add_table(rows=1, cols=2)
        tbl_ft.autofit = False
        tbl_ft.columns[0].width = Inches(3.0)
        tbl_ft.columns[1].width = Inches(3.5)
        
        # Nơi nhận - Hình 5: "Nơi nhận:" cỡ 12 Nghiêng Đậm; Danh sách cỡ 11 Đứng
        c_rec = tbl_ft.cell(0, 0).paragraphs[0]
        r_rec_head = c_rec.add_run("Nơi nhận:\n")
        r_rec_head.bold = True
        r_rec_head.italic = True
        r_rec_head.font.size = Pt(12)
        for rec in footer.get("recipients", []):
            rec_clean = rec if rec.startswith("-") else f"- {rec}"
            if not (rec_clean.endswith(";") or rec_clean.endswith(".")):
                rec_clean += ";"
            r_r = c_rec.add_run(f"{rec_clean}\n")
            r_r.font.size = Pt(11)
            
        if footer.get("author_code"):
            r_ac = c_rec.add_run(footer.get("author_code"))
            r_ac.font.size = Pt(11)
            
        # Người ký - Hình 4: Cỡ 13-14 Đứng Đậm
        c_sig = tbl_ft.cell(0, 1).paragraphs[0]
        c_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sig = footer.get("signatory", {})
        if sig.get("role"):
            r_role = c_sig.add_run(sig.get("role").upper() + "\n")
            r_role.bold = True
            r_role.font.size = Pt(14)
        if sig.get("position"):
            r_pos = c_sig.add_run(sig.get("position").upper() + "\n\n\n\n")
            r_pos.bold = True
            r_pos.font.size = Pt(14)
        if sig.get("name"):
            r_name = c_sig.add_run(sig.get("name"))
            r_name.bold = True
            r_name.font.size = Pt(14)
            
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
