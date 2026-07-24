import io
import os
import typing
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from nom.resolution.schema import NghiQuyetBase
from docxtpl import DocxTemplate

class WordExporter:
    def export(self, data: typing.Union[NghiQuyetBase, dict], template_path: str = None) -> io.BytesIO:
        if template_path and os.path.exists(template_path):
            return self._export_with_template(data, template_path)
        return self._export_default(data)
        
    def _export_with_template(self, data: typing.Union[NghiQuyetBase, dict], template_path: str) -> io.BytesIO:
        """Export using a predefined Word layout template (.docx) with Jinja2 syntax."""
        doc = DocxTemplate(template_path)
        context = data if isinstance(data, dict) else data.model_dump()
        doc.render(context)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _export_default(self, data: typing.Union[NghiQuyetBase, dict]) -> io.BytesIO:
        if isinstance(data, dict):
            # Create a mock NghiQuyetBase object to access properties, filling missing with empty string/list
            data = NghiQuyetBase(
                co_quan_ban_hanh=data.get("co_quan_ban_hanh", ""),
                so_nghi_quyet=data.get("so_nghi_quyet", ""),
                dia_danh=data.get("dia_danh", ""),
                ngay_ban_hanh=data.get("ngay_ban_hanh", ""),
                trich_yeu=data.get("trich_yeu", ""),
                can_cu_phap_ly=data.get("can_cu_phap_ly", []),
                noi_dung_dieu_khoan=data.get("noi_dung_dieu_khoan", []),
                noi_nhan=data.get("noi_nhan", []),
                chuc_vu_nguoi_ky=data.get("chuc_vu_nguoi_ky", ""),
                nguoi_ky=data.get("nguoi_ky", "")
            )
        doc = Document()
        
        # Cài đặt font mặc định
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        
        # Bảng header (Quốc hiệu tiêu ngữ)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(3)
        table.columns[1].width = Inches(3.5)
        
        cell_1 = table.cell(0, 0)
        p1 = cell_1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(data.co_quan_ban_hanh.upper())
        run1.bold = True
        run1.font.size = Pt(13)
        p1.add_run(f"\nSố: {data.so_nghi_quyet}")
        
        cell_2 = table.cell(0, 1)
        p2 = cell_2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
        run2.bold = True
        run2.font.size = Pt(13)
        p2.add_run("\nĐộc lập - Tự do - Hạnh phúc")
        
        # Ngày tháng
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_date.add_run(f"{data.dia_danh}, {data.ngay_ban_hanh}\n")
        
        # Tên loại văn bản
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run("NGHỊ QUYẾT")
        run_title.bold = True
        run_title.font.size = Pt(14)
        p_title.add_run(f"\n{data.trich_yeu}\n")
        
        # Căn cứ pháp lý
        for cc in data.can_cu_phap_ly:
            p_cc = doc.add_paragraph(f"Căn cứ {cc};")
            p_cc.paragraph_format.first_line_indent = Inches(0.5)
            
        p_quyet_nghi = doc.add_paragraph()
        p_quyet_nghi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_qn = p_quyet_nghi.add_run("\nQUYẾT NGHỊ:\n")
        run_qn.bold = True
        
        # Điều khoản
        for dieu in data.noi_dung_dieu_khoan:
            p_dieu = doc.add_paragraph(dieu)
            p_dieu.paragraph_format.first_line_indent = Inches(0.5)
            
        doc.add_paragraph("\n")
        
        # Chữ ký và nơi nhận
        table_footer = doc.add_table(rows=1, cols=2)
        
        # Nơi nhận
        cell_nn = table_footer.cell(0, 0)
        p_nn = cell_nn.paragraphs[0]
        run_nn = p_nn.add_run("Nơi nhận:\n")
        run_nn.bold = True
        run_nn.italic = True
        run_nn.font.size = Pt(12)
        for nn in data.noi_nhan:
            p_nn.add_run(f"- {nn};\n").font.size = Pt(11)
            
        # Chữ ký
        cell_ck = table_footer.cell(0, 1)
        p_ck = cell_ck.paragraphs[0]
        p_ck.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cv = p_ck.add_run(data.chuc_vu_nguoi_ky.upper())
        run_cv.bold = True
        p_ck.add_run("\n\n\n\n")
        run_ten = p_ck.add_run(data.nguoi_ky)
        run_ten.bold = True
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
